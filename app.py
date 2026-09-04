import streamlit as st
import PyPDF2
import docx
import sqlite3
import pandas as pd
import os
import re
import smtplib
from PIL import Image
from email.message import EmailMessage
from datetime import datetime
from dotenv import load_dotenv

import sys
import asyncio

# Fix for Windows asyncio Proactor socket issue in Python 3.13
if sys.platform == "win32":
    try:
        asyncio.set_event_loop_policy(asyncio.WindowsSelectorEventLoopPolicy())
    except Exception:
        pass

# --- 1. SECURITY & CONFIG ---
load_dotenv()
GROQ_API_KEY = os.getenv("GROQ_API_KEY")
GROQ_MODEL = os.getenv("GROQ_MODEL", "openai/gpt-oss-120b")
GROQ_BASE_URL = os.getenv("GROQ_BASE_URL", "https://api.groq.com/openai/v1")

# Authentication Credentials exclusively from .env
ADMIN_USERNAME = os.getenv("ADMIN_USERNAME")
ADMIN_PASSWORD = os.getenv("ADMIN_PASSWORD")

# Email SMTP Config exclusively from .env
SENDER_NAME = os.getenv("SENDER_NAME", "Recruiter")
SENDER_EMAIL = os.getenv("SENDER_EMAIL")
SENDER_PW = os.getenv("SENDER_PW")

# Initialize Groq AI Client
ai_client = None
ai_model = GROQ_MODEL

if GROQ_API_KEY and GROQ_API_KEY.strip() and not GROQ_API_KEY.startswith("your_"):
    try:
        from openai import OpenAI
        ai_client = OpenAI(
            base_url=GROQ_BASE_URL,
            api_key=GROQ_API_KEY,
        )
    except Exception as err:
        pass


logo_path = os.path.join(os.path.dirname(__file__), "logo.jpeg") if os.path.exists(os.path.join(os.path.dirname(__file__), "logo.jpeg")) else "logo.jpeg"
icon_image = Image.open(logo_path) if os.path.exists(logo_path) else None

# Set page configuration
st.set_page_config(page_title="Aditya's AI Recruiter",
                   page_icon=icon_image, layout="wide")

# --- 2. ADVANCED COLOURFUL UI STYLING ---
st.markdown("""
    <style>
    /* Main Background Gradient */
    .stApp {
     
    }
    
    /* ADDED: Border styling for Input Boxes (Username, Password) */
    .stTextInput input {
        border: 2px solid #1e3a8a !important; /* Dark blue border */
        border-radius: 8px !important;
        padding: 10px !important;
       
    }
    
    /* ADDED: Border styling for Job Description Text Area */
    .stTextArea textarea {
        border: 2px solid #28a745 !important; /* Green border */
        border-radius: 8px !important;
        padding: 10px !important;
      
    }

    /* ADDED: Border styling for File Uploader Box */
    [data-testid="stFileUploader"] {
        border: 2px solid #28a745 !important; /* Purple dashed border */
        border-radius: 8px !important;
        padding: 15px !important;
        
    }
    
    /* Premium Metric Cards */
    [data-testid="stMetricValue"] { color: #1e3a8a; font-size: 35px; font-weight: 900; }
    .stMetric { 
        
        padding: 20px; 
        border-radius: 15px; 
        box-shadow: 0 10px 20px rgba(0,0,0,0.08);
        border: 1px solid #eef2f6;
        transition: transform 0.3s ease;
    }
    .stMetric:hover { transform: translateY(-5px); }
    
    /* Specific Top Borders for Metrics to make them colorful */
    div[data-testid="column"]:nth-of-type(1) .stMetric { border-top: 5px solid #28a745; } /* Green for JD */
    div[data-testid="column"]:nth-of-type(2) .stMetric { border-top: 5px solid #6f42c1; } /* Purple for ATS */
    
    /* Colourful Section Headers */
    .section-header {
        background: linear-gradient(90deg, #1e3a8a, #00c6ff);
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
        font-size: 20px;
        font-weight: 800;
        margin-top: 25px;
        margin-bottom: 10px;
        border-bottom: 2px dashed #a0c4ff;
        padding-bottom: 5px;
    }
    
    /* Styled Buttons - Size reduced for better fit in tables */
    .stButton > button {
        background: linear-gradient(90deg, #4b6cb7 0%, #182848 100%);
        color: white;
        border: none;
        border-radius: 6px;
        padding: 5px 15px; 
        font-size: 14px;
        font-weight: bold;
        transition: all 0.3s ease;
    }
    .stButton > button:hover {
        transform: scale(1.02);
        box-shadow: 0 5px 10px rgba(0,0,0,0.2);
        color: #e0e0e0;
    }
    /* Dark & Light Mode Fix */
    input, textarea, .stTextInput input,.stTextArea textarea {
        color: var(--text-color) !important;
        -webkit-text-fill-color: var(--text-color) !important;
        background-color: transparent !important;
    }
    /* Background and Labels */
    .stApp, [data-testid="stAppViewContainer"], [data-testid="stHeader"] {
      background-color: var(--background-color) !important;
    }
    label, p, span, .stMarkdown, .stText {
      color: var(--text-color) !important;
    }
    </style>
    """, unsafe_allow_html=True)

# --- 3. DATABASE MANAGEMENT ---
DB_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'ats_pro_v4.db')


def init_db():
    conn = sqlite3.connect(DB_PATH)
    c = conn.cursor()
    c.execute('''CREATE TABLE IF NOT EXISTS candidates 
                 (id INTEGER PRIMARY KEY AUTOINCREMENT, 
                  file_name TEXT, actual_name TEXT, email TEXT, 
                  jd_score TEXT, ats_score TEXT, 
                  matched TEXT, missing TEXT, edu TEXT, 
                  exp TEXT, proj TEXT, addr TEXT, date TEXT)''')
    conn.commit()
    conn.close()


def save_to_db(file_name, actual_name, email, jd_s, ats_s, matched, missing, edu, exp, proj, addr):
    conn = sqlite3.connect(DB_PATH)
    c = conn.cursor()
    c.execute("""INSERT INTO candidates 
                 (file_name, actual_name, email, jd_score, ats_score, matched, missing, edu, exp, proj, addr, date) 
                 VALUES (?,?,?,?,?,?,?,?,?,?,?,?)""",
              (file_name, actual_name, email, jd_s, ats_s, matched, missing, edu, exp, proj, addr, datetime.now().strftime("%Y-%m-%d %H:%M")))
    conn.commit()
    conn.close()


def delete_candidate(id):
    conn = sqlite3.connect(DB_PATH)
    c = conn.cursor()
    c.execute("DELETE FROM candidates WHERE id = ?", (id,))
    conn.commit()
    conn.close()


# --- 4. ENGINE FUNCTIONS ---


def get_ai_analysis(jd, resume_text):
    """Deep contextual analysis with strict format rules and seed=42 for deterministic ATS scoring."""
    prompt = f"""
    Analyze the Resume against the JD and extract data strictly into these XML tags.
    CRITICAL RULE FOR ALL SECTIONS: DO NOT USE HTML TAGS (like <b>, <i>, <br>). ONLY use Markdown (like bold) for formatting.
    
      <thinking>
    Use this space for your rough work and step-by-step calculation.
    1. DEEP JD SCAN: Read the ENTIRE Job Description line-by-line. Pay special attention to the 'Responsibilities', 'Roles', and 'Requirements' paragraphs. Extract EVERY hidden technical skill, tool, or technology mentioned there, not just the ones in bulleted lists.
    2. Carefully check the Resume for these extracted skills (allow synonyms like 'React'='React.js', 'QA'='Testing').
    3. Count total JD skills, matched skills, and missing skills.
    4. Calculate SKILLS MATCH SCORE (out of 60 points): (Matched Skills / Total JD Skills) * 60.
    5. Evaluate EDUCATION RELEVANCE (out of 20 points): 20 if highly relevant, 10 if somewhat, 0 if irrelevant.
    6. Evaluate EXPERIENCE RELEVANCE (out of 20 points): 20 if highly relevant, 10 if somewhat, 0 if irrelevant/fresher.
    7. Sum the scores (Skills + Edu + Exp) to get the final Job Match Score.
    8. ATS EVALUATION: Evaluate the true ATS Health Score from 0-100 purely based on resume structure, formatting, and content quality. DO NOT judge based on the JD.
    Start with a base score of 100.
    MISSING SECTIONS: Check for: Contact Info, Summary, Skills, Work Experience, Education, Projects, Certifications. Deduct 5 points for EACH missing section.
    CONTENT QUALITY: Deduct 15 points if sections exist but contain NO meaningful data or dummy text.
    FORMATTING: Deduct 15 points if the text format looks garbled or messy.
    SPELLING & GRAMMAR: Deduct 10 points for spelling mistakes.
    FINAL OUTPUT FORMATE: Print ONLY the final integer number. Strictly NO text, No calculations, NO explanations. Just the percentage number(e.g.90).
    </thinking>
    
    <candidate_name>Extract the actual full name of the candidate</candidate_name>
    
    <matched>Based on the Deep JD Scan in your <thinking> section, list the skills/tools required ANYWHERE in the JD (including within the Responsibilities) that are ALSO present in the Resume. If absolutely ZERO skills match, output exactly 'None'.</matched>
    
    <missing>Based on the Deep JD Scan in your <thinking> section, list the skills/tools required ANYWHERE in the JD (including within the Responsibilities) but completely MISSING from the Resume. If all required skills are present, output exactly 'None'.</missing>
    
    <jd_score>Output ONLY the final calculated integer total score from the <thinking> phase. </jd_score>
    
    <ats_score>
    Output ONLY the final calculated integer total score from the <thinking> phase. Print ONLY the final integer number. Strictly NO text, No calculations, NO explanations. Just the percentage number(e.g.90). 
    </ats_score>
    
    <education>
    Extract clear and clean education details. Format like this: Degree/Major from University/College Name (Year).
    </education>
    
    <experience>
    Format each experience exactly like this:
    - **Job Title at Company Name** (Employment type | Duration)
    Write 1 or 2 clear sentences explaining their main responsibilities. Output 'No valid experience found' ONLY if the candidate is a pure fresher with absolutely zero experience written. 
    </experience>
    
    <projects>
    Format each project strictly using bullet points:
    - **[Project Name]**: Short, 1-line summary of what it is and the technologies used.
    (Output 'Not Applicable' if empty).
    </projects>
    
    <address>
    Extract the full address. If not found, strictly output 'Not Provided'. DO NOT invent details.
    </address>
   
    JD: {jd}
    Resume: {resume_text}
    """
    if not ai_client:
        raise ValueError("No AI API client configured! Please check GROQ_API_KEY in your .env file.")

    response = ai_client.chat.completions.create(
        model=ai_model,
        messages=[{"role": "user", "content": prompt}],
        temperature=0.0,
    )
    return response.choices[0].message.content



def send_invitation_email(target, actual_name):
    """SMTP Email Delivery using actual candidate name."""
    try:
        msg = EmailMessage()
        msg['Subject'] = " Interview Invitation: Next Steps for Your Application "
        msg['From'] = f"{SENDER_NAME} <{SENDER_EMAIL}>"
        msg['To'] = target
        msg.set_content(f"Dear {actual_name},\n\nThank you for your interest in joining our team. We have carefully reviewed your application and resume. We are pleased to inform you that your profile has been shortlisted for the next stage of our selection process. \n\nAs the next step,we would like to schedule an interview with you to discuss your past projects,experience,and the role in more detail. We aim to schedule this interview within the next 3 days,please stay connected!\n\nBest Regards,\nHR Team")

        with smtplib.SMTP_SSL('smtp.gmail.com', 465) as smtp:
            smtp.login(SENDER_EMAIL, SENDER_PW)
            smtp.send_message(msg)
        return True
    except Exception as e:
        st.error(f"Email Error Details: {e}")
        return False


# --- 5. APP STATE INITIALIZATION ---
if 'auth' not in st.session_state:
    st.session_state['auth'] = False
if 'results' not in st.session_state:
    st.session_state['results'] = []
if 'jd_input' not in st.session_state:
    st.session_state['jd_input'] = ""
if 'view_record_id' not in st.session_state:
    st.session_state['view_record_id'] = None

# --- 6. USER INTERFACE COMPONENTS ---


def login_screen():
    st.markdown("<br><br><br>", unsafe_allow_html=True)
    col1, col2, col3 = st.columns([1, 2, 1])
    with col2:
        st.markdown(
            "<h1 style='text-align: center; color: #1e3a8a;'>🔒 Recruiter Portal</h1>", unsafe_allow_html=True)
        # Bounded by CSS class .stTextInput
        user = st.text_input("**Username**")
        pw = st.text_input("**Password**", type="password")
        if st.button("Access Dashboard", use_container_width=True):
            if not ADMIN_USERNAME or not ADMIN_PASSWORD:
                st.error("❌ Portal credentials are not configured. Please set ADMIN_USERNAME and ADMIN_PASSWORD in your .env file.")
            elif user == ADMIN_USERNAME and pw == ADMIN_PASSWORD:
                st.session_state['auth'] = True
                st.rerun()
            else:
                st.error("Invalid Username/Password")


def main_dashboard():
    st.sidebar.markdown("## ⚙️ **Control Panel**")
    choice = st.sidebar.radio(
        "**Navigation**", ["Smart Dashboard", "Database History"])

    if st.sidebar.button("Logout"):
        st.session_state['auth'] = False
        st.rerun()

    if choice == "Smart Dashboard":
        st.markdown(
            "<h1 style='color: #1e3a8a;'>Recruitment Process</h1>", unsafe_allow_html=True)

        # Bounded by CSS class .stTextArea
        jd_input = st.text_area("**Step 1: Paste Job Description**",
                                value=st.session_state['jd_input'], height=150, placeholder="Required: Job Description...")
        st.session_state['jd_input'] = jd_input

        # Bounded by CSS targeting data-testid="stFileUploader"
        uploaded_files = st.file_uploader("**Step 2: Upload Resumes (PDF,DOCX,TXT)**", type=[
                                          "pdf", "docx", "txt"], accept_multiple_files=True)

        if st.button("Click Here For Analyze & Rank Resumes"):
            if st.session_state['jd_input'] and uploaded_files:
                st.session_state['results'] = []

                with st.spinner(" Analyzing Resumes... Please wait."):
                    for file in uploaded_files:
                        resume_text = ""
                        file_ext = file.name.split('.')[-1].lower()
                        # 1. PDF Ectract Logic
                        if file_ext == "pdf":
                            pdf_reader = PyPDF2.PdfReader(file)
                            resume_text = "".join(
                                [p.extract_text() for p in pdf_reader.pages])
                        # 2. DOCX Extract Logic
                        elif file_ext == "docx":
                            doc = docx.Document(file)
                            doc_parts = [para.text for para in doc.paragraphs if para.text.strip()]
                            for tbl in doc.tables:
                                for row in tbl.rows:
                                    row_text = " | ".join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                                    if row_text:
                                        doc_parts.append(row_text)
                            resume_text = "\n".join(doc_parts)
                        # 3. TXT Extract Logic
                        elif file_ext == "txt":
                            resume_text = file.read().decode("utf-8")

                        if resume_text.strip():
                            try:
                                raw_analysis = get_ai_analysis(
                                    st.session_state['jd_input'], resume_text)
                            except Exception as api_err:
                                err_str = str(api_err)
                                if "429" in err_str or "rate limit" in err_str.lower():
                                    st.error(f"❌ Groq Rate Limit Exceeded for {file.name}: Rate limit reached. Please check your Groq quota or retry in a few moments.")
                                else:
                                    st.error(f"❌ AI Error analyzing {file.name}: {err_str}")
                                continue
                        else:
                            st.warning(f"⚠️ Could not extract text from {file.name}")
                            continue

                        def parse_tag(tag, source):
                            match = re.search(
                                f"<{tag}>(.*?)</{tag}>", source, re.S)
                            return match.group(1).strip() if match else "N/A"

                        email_find = re.findall(
                            r'[a-zA-Z0-9_.+-]+@[a-zA-Z0-9-]+\.[a-zA-Z0-9-.]+', resume_text)
                        candidate_email = email_find[0] if email_find else "Not Found"

                        parsed_res = {
                            "file_name": file.name,
                            "actual_name": parse_tag("candidate_name", raw_analysis),
                            "email": candidate_email,
                            "jd_s": parse_tag("jd_score", raw_analysis),
                            "ats_s": parse_tag("ats_score", raw_analysis),
                            "matched": parse_tag("matched", raw_analysis),
                            "missing": parse_tag("missing", raw_analysis),
                            "edu": parse_tag("education", raw_analysis),
                            "exp": parse_tag("experience", raw_analysis),
                            "proj": parse_tag("projects", raw_analysis),
                            "addr": parse_tag("address", raw_analysis)
                        }

                        save_to_db(
                            parsed_res['file_name'], parsed_res['actual_name'], parsed_res['email'],
                            parsed_res['jd_s'], parsed_res['ats_s'], parsed_res['matched'],
                            parsed_res['missing'], parsed_res['edu'], parsed_res['exp'],
                            parsed_res['proj'], parsed_res['addr']
                        )
                        st.session_state['results'].append(parsed_res)

                def get_numeric_score(res):
                    nums = re.findall(r'\d+', str(res.get('jd_s', '0')))
                    return int(nums[0]) if nums else 0

                st.session_state['results'].sort(
                    key=get_numeric_score, reverse=True)
                st.success("✅ Analysis Complete & Ranked by Best Match!")
            else:
                st.warning(
                    "⚠️ Please provide both Job Description and Resumes.")

        for r in st.session_state['results']:
            with st.expander(f"📊 Candidate Report: {r['file_name']} | Candidate Name: {r['actual_name']}", expanded=True):
                col1, col2 = st.columns(2)
                col1.metric("JD Match Score", f"{r['jd_s']}%")
                col2.metric("ATS Score", f"{r['ats_s']}%")

                st.markdown(
                    "<div class='section-header'>✅ Matched Skills</div>", unsafe_allow_html=True)
                st.write(r['matched'])

                st.markdown(
                    "<div class='section-header'>❌ Missing Skills</div>", unsafe_allow_html=True)
                st.write(r['missing'])

                st.markdown(
                    "<div class='section-header'>🎓 Education</div>", unsafe_allow_html=True)
                st.markdown(r['edu'])

                st.markdown(
                    "<div class='section-header'>💼 Experience</div>", unsafe_allow_html=True)
                st.markdown(r['exp'])

                st.markdown(
                    "<div class='section-header'>🏗️ Projects</div>", unsafe_allow_html=True)
                st.markdown(r['proj'])

                st.markdown(
                    "<div class='section-header'>📍 Candidate Address</div>", unsafe_allow_html=True)
                st.write(r['addr'])

                st.markdown("<br>", unsafe_allow_html=True)

                btn_text = f"📧 Send Interview Invite to {r['actual_name'] if r['actual_name'] != 'N/A' else r['file_name']}"
                if st.button(btn_text, key=f"mail_{r['file_name']}"):
                    if r['email'] != "Not Found":
                        with st.spinner("Dispatching Email..."):
                            name_to_use = r['actual_name'] if r['actual_name'] != 'N/A' else 'Candidate'
                            if send_invitation_email(r['email'], name_to_use):
                                st.success(
                                    f"Invite successfully sent to {r['email']}")
                            else:
                                st.error(
                                    "Email failed. Please check App Password or Port Settings.")
                    else:
                        st.error("Email not found in resume.")

    elif choice == "Database History":
        st.markdown(
            "<h1 style='color: #1e3a8a;'>📁 Candidate Records</h1>", unsafe_allow_html=True)
        conn = sqlite3.connect(DB_PATH)
        df = pd.read_sql_query(
            "SELECT * FROM candidates ORDER BY id ASC", conn)
        conn.close()

        if not df.empty:
            h1, h2, h3, h4, h5, h6 = st.columns([0.8, 2, 2.5, 2, 1.5, 1])
            h1.markdown("S.No")
            h2.markdown("File Name")
            h3.markdown("Email")
            h4.markdown("Scores")
            h5.markdown("Report")
            h6.markdown("Action")
            st.markdown("---")

            for idx, (index, row) in enumerate(df.iterrows(), start=1):
                c1, c2, c3, c4, c5, c6 = st.columns([0.8, 2, 2.5, 2, 1.5, 1])
                c1.write(idx)
                c2.write(f"{row['file_name']}")
                c3.write(row['email'])
                c4.write(f"JD: {row['jd_score']}% | ATS: {row['ats_score']}%")

                if c5.button("📄 View", key=f"view_{row['id']}"):
                    if st.session_state['view_record_id'] == row['id']:
                        st.session_state['view_record_id'] = None
                    else:
                        st.session_state['view_record_id'] = row['id']

                if c6.button("🗑️", key=f"del_{row['id']}"):
                    delete_candidate(row['id'])
                    st.rerun()

                if st.session_state['view_record_id'] == row['id']:
                    st.markdown("---")
                    st.markdown(
                        f"<h3 style='color: #1e3a8a;'>📋 Detailed Report: {row['actual_name']}</h3>", unsafe_allow_html=True)
                    st.markdown(f"✅ **Matched Skills:** {row['matched']}")
                    st.markdown(f"❌ **Missing Skills:** {row['missing']}")
                    st.markdown("<br>", unsafe_allow_html=True)
                    st.markdown("🎓 **Education:**")
                    st.markdown(row['edu'])
                    st.markdown("💼 **Experience:**")
                    st.markdown(row['exp'])
                    st.markdown("🏗️ **Projects:**")
                    st.markdown(row['proj'])
                    st.markdown("<br>", unsafe_allow_html=True)
                    st.markdown(f"📍 **Candidate Address:** {row['addr']}")
                    st.markdown("---")
        else:
            st.info("No records found in database.")


# --- 7. START APPLICATION ---
init_db()
if not st.session_state['auth']:
    login_screen()
else:
    main_dashboard()
